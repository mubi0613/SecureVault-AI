import json
import os
import base64
import hashlib
import time
from datetime import datetime
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from textblob import TextBlob
from fpdf import FPDF
from docx import Document
import io
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import secrets
import string
from google import genai  # Gemini
from dotenv import load_dotenv
import nltk

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab')

# This line reads the .env file and makes the API key available
load_dotenv(override=True)  # 'override=True' forces it to refresh the key if you changed .env

NOTES_FILE = "notes.json"
CONFIG_FILE = "vault_config.json"

# Load a small, fast model for embeddings (runs locally)
embed_model = SentenceTransformer('all-MiniLM-L6-v2')

# --- 0. PIN & KEY LOGIC ---
def get_pin_hash(pin: str):
    """Creates a secure SHA-256 hash of the PIN"""
    return hashlib.sha256(pin.encode()).hexdigest()

def is_vault_initialized():
    """Checks if a user has set a PIN yet"""
    return os.path.exists(CONFIG_FILE)

def generate_recovery_key():
    """Generates a random 16-character recovery string"""
    chars = string.ascii_uppercase + string.digits
    return ''.join(secrets.choice(chars) for _ in range(16))

def initialize_vault(pin: str, recovery_key: str):
    """Saves the user's initial PIN hash and recovery key hash"""
    with open(CONFIG_FILE, "w") as f:
        json.dump({"pin_hash": get_pin_hash(pin),
                   "recovery_hash": get_pin_hash(recovery_key) # We hash the recovery key too!
                   }, f)

def verify_recovery_key(input_key: str):
    """Verifies the recovery key against the stored hash"""
    if not is_vault_initialized():
        return False
    with open(CONFIG_FILE, "r") as f:
        stored_recovery_hash = json.load(f).get("recovery_hash")
    return get_pin_hash(input_key) == stored_recovery_hash

def verify_pin(input_pin: str):
    """Verifies input against stored hash"""
    if not is_vault_initialized():
        return False
    with open(CONFIG_FILE, "r") as f:
        stored_hash = json.load(f).get("pin_hash")
    return get_pin_hash(input_pin) == stored_hash

# --- 1. KEY GENERATION (From your PIN) ---
def generate_key(pin: str):
    """Derives a functional encryption key from the user's PIN"""
    password = pin.encode()
    # We use a static salt for this local project (Professional apps use random salts)
    salt = b'stable_salt_123' 
    kdf = PBKDF2HMAC(
        algorithm=hashes.SHA256(),
        length=32,
        salt=salt,
        iterations=100000,
    )
    key = base64.urlsafe_b64encode(kdf.derive(password))
    return Fernet(key)

# --- 2. ENCRYPTION / DECRYPTION ---
def encrypt_data(data_string, pin):
    """Turns readable text into scrambled code"""
    f = generate_key(pin)
    return f.encrypt(data_string.encode()).decode()

def decrypt_data(encrypted_string, pin):
    """Turns scrambled code back into readable text"""
    try:
        f = generate_key(pin)
        return f.decrypt(encrypted_string.encode()).decode()
    except Exception:
        return "[Decryption Error: Check PIN]"

# --- 3. UPDATED LOAD/SAVE ---
def load_notes():
    if os.path.exists(NOTES_FILE):
        with open(NOTES_FILE, "r") as f:
            return json.load(f)
    return []

def save_notes(notes_list):
    with open(NOTES_FILE, "w") as f:
        json.dump(notes_list, f, indent=4)

def get_page_heading(is_unlocked):
    return "🛡️ Safe Vault" if is_unlocked else "📝 My Notes"

def get_filtered_notes(all_notes, is_unlocked, search_query):
    visible = all_notes
    if not is_unlocked:
        visible = [n for n in all_notes if not n.get('secret', False)]
    
    query = search_query.lower()
    return [n for n in visible if query in n['title'].lower() or query in n['content'].lower()]

# --- 4. AI FEATURES ---
def ai_summarize_text(text):
    """Uses NLP to extract key points from long notes."""
    if len(text) < 50:
        return text # Don't summarize very short notes
    
    blob = TextBlob(text)
    # Extracting sentences and picking the top 2 for a summary
    sentences = [str(s) for s in blob.sentences]
    
    if len(sentences) > 2:
        summary = "✨ AI Summary:\n" + "\n".join([f"• {s}" for s in sentences[:2]])
        
        # NEW: Log the usage here!
        track_usage(text, type="input") 
        track_usage(summary, type="output")
        return summary
    return text

# --- 5. EXPORTS ---
# The Logic Functions : These handle the actual file creation.
def create_pdf(title, content):
    try:
        pdf = FPDF()
        pdf.add_page()

        # Updated to match your folder: fonts -> dejavu-sans -> DejaVuSans.ttf
        font_path = os.path.join("fonts", "dejavu-sans", "DejaVuSans.ttf")
        
        if os.path.exists(font_path):
            pdf.add_font('DejaVu', '', font_path, uni=True)
            pdf.set_font('DejaVu', '', 14)
        else:
            # This is where it was falling back before, causing the "?" symbols
            pdf.set_font("Arial", "B", 16)

        # Write Title
        pdf.multi_cell(0, 10, title)
        pdf.ln(5)
        
        # Write Content
        if os.path.exists(font_path):
            pdf.set_font('DejaVu', '', 12)
        else:
            pdf.set_font("Arial", "", 12)
            
        pdf.multi_cell(0, 10, content)
        
        # Output as bytes
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        return f"PDF Error: {str(e)}".encode('latin-1')
    
def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 6. AI & RESOURCE TRACKER ---
def track_usage(text, type="input"):
    # Rough estimation: 1 token approx 4 characters
    tokens = len(text) / 4
    cost = (tokens / 1000) * 0.000125 # Estimate for Gemini 1.5 Flash
    
    # Save to a local JSON file to persist data
    usage_file = "usage_stats.json"
    stats = {"total_tokens": 0, "total_cost": 0.0}
    
    if os.path.exists(usage_file):
        with open(usage_file, "r") as f:
            stats = json.load(f)
            
    stats["total_tokens"] += tokens
    stats["total_cost"] += cost
    
    with open(usage_file, "w") as f:
        json.dump(stats, f)
    
    return stats

# --- 7. RAG & FEEDBACK ---
def create_vector_index(notes):
    """Turns notes into a searchable mathematical index"""
    if not notes:
        return None, []
    
    # Extract only the text content
    text_data = [n['content'] for n in notes]
    
    # Convert text to vectors (embeddings)
    embeddings = embed_model.encode(text_data)
    
    # Create the FAISS index
    dimension = embeddings.shape[1]
    index = faiss.IndexFlatL2(dimension)
    index.add(np.array(embeddings).astype('float32'))
    
    return index, text_data

def query_vault(query, index, text_data, top_k=2):
    """Finds the most relevant notes for a question"""
    if index is None:
        return "No notes found to search."
        
    # Convert question to vector
    query_vector = embed_model.encode([query])
    
    # Search the index
    distances, indices = index.search(np.array(query_vector).astype('float32'), top_k)
    
    # Pull the relevant text chunks
    results = [text_data[i] for i in indices[0] if i != -1]
    return results

def log_feedback(query, answer, context, status):
    log_file = "feedback_log.json"
    entry = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "query": query,
        "answer": answer,
        "context_used": context,
        "status": status  # "Correct" or "Wrong"
    }
    
    logs = []
    if os.path.exists(log_file):
        with open(log_file, "r") as f:
            logs = json.load(f)
            
    logs.append(entry)
    with open(log_file, "w") as f:
        json.dump(logs, f, indent=4)

# --- 8. GEMINI AI ENGINE ---
def get_gemini_response(user_query, context_str):
    """Connects to Google Gemini API for free AI logic"""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        return "⚠️ Error: GEMINI_API_KEY not found in Hugging Face Secrets."
    
    try:
        # New initialization style
        client = genai.Client(api_key=api_key)
        
        prompt = f"""
        You are a secure vault assistant. Use the following retrieved notes to answer.
        Notes Context: {context_str}
        User Question: {user_query}
        """
        
        # Inside your loop or before you call the model
        time.sleep(2) # Wait 2 seconds between requests
        # Simplified model call
        response = client.models.generate_content(
            model="gemini-2.5-flash-lite", 
            contents=prompt
        )
        
        track_usage(prompt, type="input")
        track_usage(response.text, type="output")
        return response.text
    except Exception as e:
        return f"❌ AI Engine Error: {str(e)}"